import io
import os
import re
import sys
from typing import Union
from src.entity.artifact_entity import ParsedDocument
from src.exception.custom_exception import CustomException
from src.logger.logging import logger

try:
    import pymupdf as fitz
except ImportError:
    try:
        import fitz
    except ImportError:
        fitz = None

try:
    from docx import Document
except ImportError:
    Document = None


class DocumentParser:
    """
    Component for parsing and extracting clean textual data from multi-format resume documents (PDF, DOCX, TXT).
    """

    def __init__(self):
        logger.info("Initializing DocumentParser component")

    @staticmethod
    def _clean_text(raw_text: str) -> str:
        if not raw_text:
            return ""
        # Remove non-standard unicode bullets and dashes
        cleaned = re.sub(r'[\u2022\u2023\u25E6\u2043\u2219\u00B7\u25AA\u25AB\u2013\u2014]', ' ', raw_text)
        cleaned = cleaned.replace('\xa0', ' ').replace('\r', '\n')
        # Consolidate excessive whitespace
        cleaned = re.sub(r'\n\s*\n+', '\n\n', cleaned)
        cleaned = re.sub(r'[ \t]+', ' ', cleaned)
        return cleaned.strip()

    def parse_pdf_bytes(self, stream_bytes: bytes) -> str:
        if fitz is None:
            raise ImportError("PyMuPDF is required for PDF parsing.")
        
        extracted_pages = []
        try:
            with fitz.open(stream=stream_bytes, filetype="pdf") as doc:
                for page_idx in range(len(doc)):
                    page = doc.load_page(page_idx)
                    page_text = page.get_text("text")
                    if page_text:
                        extracted_pages.append(page_text)
            return "\n".join(extracted_pages)
        except Exception as e:
            logger.error(f"Error parsing PDF bytes: {str(e)}")
            raise CustomException(e, sys)

    def parse_docx_bytes(self, stream_bytes: bytes) -> str:
        if Document is None:
            raise ImportError("python-docx is required for DOCX parsing.")

        extracted_lines = []
        try:
            doc = Document(io.BytesIO(stream_bytes))
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_lines.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        extracted_lines.append(" | ".join(row_cells))
            return "\n".join(extracted_lines)
        except Exception as e:
            logger.error(f"Error parsing DOCX bytes: {str(e)}")
            raise CustomException(e, sys)

    def parse_txt_bytes(self, stream_bytes: bytes) -> str:
        try:
            try:
                return stream_bytes.decode("utf-8")
            except UnicodeDecodeError:
                return stream_bytes.decode("latin-1", errors="ignore")
        except Exception as e:
            logger.error(f"Error decoding TXT bytes: {str(e)}")
            raise CustomException(e, sys)

    def extract(self, file_source: Union[bytes, io.BytesIO, str], filename: str) -> ParsedDocument:
        """
        Parses the input file source into a structured ParsedDocument artifact.
        """
        try:
            filename_lower = filename.lower()
            
            # Check if file_source is already extracted plain text string
            if isinstance(file_source, str) and not os.path.exists(file_source):
                cleaned_text = self._clean_text(file_source)
                words = cleaned_text.split()
                return ParsedDocument(
                    filename=filename,
                    raw_text=cleaned_text,
                    word_count=len(words),
                    char_count=len(cleaned_text)
                )

            # Read bytes from various input sources
            if hasattr(file_source, "read"):
                raw_bytes = file_source.read()
                if hasattr(file_source, "seek"):
                    file_source.seek(0)
            elif isinstance(file_source, io.BytesIO):
                raw_bytes = file_source.getvalue()
            elif isinstance(file_source, bytes):
                raw_bytes = file_source
            elif isinstance(file_source, str) and os.path.exists(file_source):
                with open(file_source, "rb") as f:
                    raw_bytes = f.read()
            else:
                raw_bytes = str(file_source).encode("utf-8")

            if filename_lower.endswith(".pdf"):
                raw_text = self.parse_pdf_bytes(raw_bytes)
            elif filename_lower.endswith(".docx"):
                raw_text = self.parse_docx_bytes(raw_bytes)
            elif filename_lower.endswith(".txt"):
                raw_text = self.parse_txt_bytes(raw_bytes)
            else:
                raw_text = self.parse_txt_bytes(raw_bytes)

            cleaned_text = self._clean_text(raw_text)
            words = cleaned_text.split()

            return ParsedDocument(
                filename=filename,
                raw_text=cleaned_text,
                word_count=len(words),
                char_count=len(cleaned_text)
            )
        except Exception as e:
            logger.error(f"Failed to process document {filename}: {str(e)}")
            raise CustomException(e, sys)
